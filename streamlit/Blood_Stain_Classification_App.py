import streamlit as st
import os
import sys
import numpy as np
import pandas as pd
import torch
import cv2
from PIL import Image
from torchvision import transforms
from torch.nn.functional import softmax
from os.path import dirname as up
from io import BytesIO
import uuid
from streamlit_option_menu import option_menu
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import csv

sys.path.append(up(up(os.path.abspath(__file__))))
from utils import utils
from src.dataloader.labels import LABELS
from config.utils import load_config
from src.model import finetune_resnet
from src.gradcam import GradCam



def load_definitions(csv_path):
    """
    Load class definitions from a CSV file.
    """
    definitions = {}
    with open(csv_path, mode='r', encoding='utf-8') as file:
        reader = csv.reader(file, delimiter=',', quotechar='"')
        for row in reader:
            if len(row) == 2:
                class_name, definition = row
                definitions[class_name.strip()] = definition.strip()
    return definitions

definitions = load_definitions("streamlit/terminology.csv")



def generate_word_document():
    """
    Generate a Word document containing analysis results.
    """
    doc = Document()
    doc.add_heading('Blood Stain Classification Report', 0)

    for idx, row in st.session_state["results_df"].iterrows():
        doc.add_heading(f"Image: {row['Image Name']}", level=1)
        doc.add_paragraph(f"Predicted Label: {row['Predicted Label']}")
        doc.add_paragraph(f"Definition: {row['Definition']}")

        img_path = os.path.join("tempDir", row['Image Name'])
        saliency_path = os.path.join("tempDir", "saliency_map_"+row['Image Name'])

        if img_path:
            doc.add_heading('Original Image', level=2)
            doc.add_picture(img_path, width=Inches(2.0))

        if saliency_path:
            doc.add_heading('Saliency Map', level=2)
            doc.add_picture(saliency_path, width=Inches(2.0))
        doc.add_page_break()

    path = os.path.join(st.session_state["output_path"], f'{st.session_state["file_name"]}.docx')
    doc.save(path)
    st.success(f"Word document saved at {path}")

def save_uploaded_file(uploadedfile, file_name):
    """
    Save an uploaded file to a temporary directory.
    """
    temp_dir = "tempDir"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    file_path = os.path.join(temp_dir, file_name)
    with open(file_path, "wb") as f:
        f.write(uploadedfile.getbuffer())

def save_saliency_map(saliency_map, file_name):
    """
    Save a saliency map as an image in a temporary directory.
    """
    temp_dir = "tempDir"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    file_path = os.path.join(temp_dir, file_name)
    cv2.imwrite(file_path, saliency_map * 255)
    return file_path

def display_home_page():
    """
    Display the home page of the application.
    """
    st.title("Blood Stain Classification")
    st.write("")
    st.write("")

    if "results_df" not in st.session_state:
        st.session_state["results_df"] = pd.DataFrame(
            columns=['Image Name', 'Predicted Label'])

    if "image_files" not in st.session_state:
        st.session_state["image_files"] = None

    if "predict_button" not in st.session_state:
        st.session_state["predict_button"] = False

    image_files = st.file_uploader("Upload Image",
                                   type=['png', 'jpg', 'jpeg', 'PNG', 'JPG', 'JPEG'],
                                   accept_multiple_files=True)

    st.write("")
    if image_files:
        st.session_state["image_files"] = image_files
        for image_file in image_files:
            save_uploaded_file(image_file, image_file.name)

    if st.session_state["image_files"]:
        display_uploaded_images()

    st.divider()
    col1, col2, col3, col4 = st.columns(4)


    prediction_type = col1.radio("Results format :", ["Full Classification", "Prediction Only"], key="prediction_type")
    if "prediction_type" not in st.session_state:
        st.session_state["prediction_type"] = prediction_type

    col4.write("")
    col4.write("")
    if col4.button("Predict"):
        st.session_state["predict_button"] = True

    if st.session_state["predict_button"] and image_files:
        predict_images(st.session_state["prediction_type"])

    output_path = col2.text_input('Path to save results', value=os.getcwd(), key="OP")
    st.session_state["output_path"] = output_path
    file_name = col3.text_input('File name', value='results', key="FN")
    st.session_state["file_name"] = file_name


def display_uploaded_images():
    col0, col1, col2, col3 = st.columns(4)

    for i, image_file in enumerate(st.session_state["image_files"]):
        image = Image.open(image_file)
        col_num = i % 4
        col = [col0, col1, col2, col3][col_num]
        col.image(image, caption=image_file.name, width=200)

@st.cache_resource
def load(config_file_name):
    config = load_config(config_file_name)
    device = utils.get_device(device_config="cuda:0")
    config_in_log_dir = os.path.dirname(config_file_name)
    model = finetune_resnet.get_finetuneresnet(config)
    weight = utils.load_weights(config_in_log_dir, device=device)
    model.load_dict_learnable_parameters(state_dict=weight, strict=True)
    model = model.to(device)
    model.eval()
    transform = transforms.Compose([transforms.ToTensor(), transforms.Resize((256, 256))])
    gradcam = GradCam(model=model)
    return config, device, model, transform, gradcam


class choose_class():
    def __init__(self, question, top_options, other_options, key):
        self.question = question
        self.top_options = top_options + ["Autre"]
        self.other_options = other_options
        self.key = key

    def save_modification(self):
        return self.selected_class

    def choosing(self):
        with st.form(key='form' + str(self.key)):
            st.write(self.question)
            chosen_class = st.radio('', self.top_options, key=self.key)
            if chosen_class == "Autre":
                self.selected_class = st.selectbox('Choose from other classes:', self.other_options,
                                                   key=f"other_{self.key}")
            else:
                self.selected_class = chosen_class
            st.form_submit_button('Submit', on_click=self.save_modification)


def predict_images(type="Full Classification"):
    """
    Predict labels for uploaded images.
    """
    temperature = 1.5
    config_file = os.path.join('logs_previous_models', 'retrain_resnet_allw_img256_2', 'config.yaml')

    if st.session_state["image_files"]:
        config, device, model, transform, gradcam = load(config_file)
    if type == "Full Classification":
            for i, image_file in enumerate(st.session_state["image_files"]):
                st.subheader(f"Prediction of image {image_file.name}")
                st.write("Plotting the probability distribution of the classes...")
                model = model.to(device)

                image = Image.open(image_file)
                image_tensor = transform(image).unsqueeze(0).to(device)
                with torch.no_grad():
                    y_pred = model.forward(image_tensor)
                y_pred = softmax(y_pred / temperature, dim=-1)
                numpy_y_pred_prob = y_pred.cpu().numpy()[0].reshape(1, -1)

                def extract_number(label):
                    """
                    Extract the numeric part from the label for sorting.
                    """
                    return int(label.split('-')[0])

                sorted_labels = sorted(LABELS, key=extract_number)

                y_pred_df = pd.DataFrame(numpy_y_pred_prob, columns=LABELS)
                y_pred_df = y_pred_df.T
                y_pred_df = y_pred_df.reindex(sorted_labels)

                y_pred_df.index = pd.CategoricalIndex(y_pred_df.index, categories=sorted_labels, ordered=True)
                y_pred_df = y_pred_df.sort_index()


                col1, col2, col3 = st.columns(3)

                np_image = image_tensor.cpu().numpy().squeeze().transpose(1, 2, 0)
                col1.image(np_image, caption='Original Image', use_column_width=True, width=300)

                visualizations = gradcam.forward(image_tensor)
                saliency_map = visualizations[0]
                saliency_map = saliency_map / 255
                saliency_path = save_saliency_map(saliency_map, f'saliency_map_{image_file.name}')
                col2.image(saliency_path, caption='Saliency Map', use_column_width=True, width=300)


                col3.bar_chart(y_pred_df, width=200)
                top_classes = np.argsort(-y_pred.cpu().numpy()[0])[:3]
                top_classes_labels = [LABELS[idx] for idx in top_classes]
                res = LABELS[y_pred.argmax(dim=-1).item()]

                definition = definitions.get(res, "No definition available")
                new_class = choose_class('Choose the best class :', top_classes_labels, sorted_labels, key=i)
                new_class.choosing()
                res = new_class.save_modification()
                definition = definitions.get(res, "No definition available")

                if image_file.name in st.session_state["results_df"]['Image Name'].values:
                    index = st.session_state["results_df"][st.session_state["results_df"]['Image Name'] == image_file.name].index[0]
                    st.session_state["results_df"].loc[index, 'Predicted Label'] = res
                    st.session_state["results_df"].loc[index, 'Definition'] = definition
                else:
                    new_row = pd.DataFrame({
                        'Image Name': [image_file.name],
                        'Predicted Label': [res],
                        'Definition': [definition]
                    })
                    st.session_state["results_df"] = pd.concat([st.session_state["results_df"], new_row], ignore_index=True)

                col1.success(f"**This image is classified as:** {res}")
    if type == "Prediction Only":
        pass

    st.dataframe(st.session_state["results_df"])
    docx_save_button = st.button("Download as Word document")
    if docx_save_button:
        generate_word_document()

def display_user_guide():
    """
    Display the user guide for the application.
    """
    st.write("""
# Manuel d'Utilisation de l'Application "Blood Stain Classification"

Bienvenue dans l'application "Blood Stain Classification". Cette application a été conçue pour aider à l'analyse des traces de sang en utilisant l'Intelligence Artificielle. Ce manuel vous guidera à travers les étapes nécessaires pour utiliser efficacement l'application.

## 1. Téléchargement des Images

La première étape consiste à télécharger les images que vous souhaitez analyser. Assurez-vous que les images sont au format PNG, JPG ou JPEG.

## 2. Paramètres

#### Type de prédiction
Vous devez sélectionner le type de prédiction que vous souhaitez obtenir. Deux options sont disponibles :
- **Prédictions Seulement :** Cette option vous permet d'obtenir les prédictions de chaque image sous forme de tableau au format CSV seulement.
- **Analyse Complète :** Cette option vous offre une analyse détaillée pour chaque image, y compris les trois meilleures classes prédites, ainsi que la carte de salience.

#### Sauvegarde 
Vous devez également indiquer le chemin d'accès à un dossier existant dans lequel vous voulez enregistrer les résultats de prédiction et donner un nom au fichier qui sera produit. 

## 3. Analyse des Images

Une fois les options sélectionnées, cliquez sur le bouton "Predict" pour lancer l'analyse des images. Vous pourrez alors visualiser les résultats de l'analyse pour chaque image téléchargée.

## 4. Téléchargement des Résultats

Après avoir terminé l'analyse, vous avez la possibilité de télécharger les résultats sous forme de document Word ou de document CSV, selon le type de prédiction sélectionné.

## 5. Important
- Vous devez vous assurer de bien raffraichir la page avant de faire de nouvelles prédictions ou changer de mode
- Veuillez indiquer le bon chemin d'accès pour l'enregistrement des résultats

## 5. À Propos de Nous

Pour en savoir plus sur l'application et son développement, consultez la section "À Propos de Nous".
""")

def display_about_us():
    """
    Display information about the application and its developers.
    """
    st.write("""L'application Blood Stain Classification a été développée par le Laboratoire d'Analyse Criminalistique (LAC) pour aider les professionnels de la criminalistique dans l'analyse de traces de sang. 
    Son développement a été réalisé en collaboration avec l'équipe de recherche Qarma du Laboratoire d'Informatique et Système de Marseille, et a été développé par Cléa Han, Yanis Labeyrie Adrien Zabban, Nesrine Ghannay et Clara KARKACH supervisé par M. Stéphane AYACHE et M. Ronan SICRE.
    Pour toute question ou suggestion, veuillez nous contacter à nesrine.ghannay@etu.univ-amu.fr.""")


def main():
    """
    Main function to run the Blood Stain Classification application.
    """
    st.set_page_config(page_title="Blood Stain Classification", page_icon=":drop_of_blood:", layout="wide")
    st.image("streamlit/ressources/lac_logo.jpg", width=150)

    menu = option_menu(None, ["Home", "User Guide", "About Us"],
                       icons=['house', 'list-task', ""],
                       menu_icon="cast", default_index=0, orientation="horizontal")

    if menu == "Home":
        display_home_page()
    elif menu == "User Guide":
        display_user_guide()
    elif menu == "About Us":
        display_about_us()


if __name__ == '__main__':
    main()